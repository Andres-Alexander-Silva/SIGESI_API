"""Tests de los endpoints de formatos para administradores y directores de semillero.

Cubre /api/v1/informes/formularios-docente/ (paquete .zip según rol/tipo de
vinculación) y /api/v1/informes/formularios-docente/descargar/?form_name=<slug>
(formato individual): camino feliz por rol permitido, 403 por rol no permitido, y
las reglas de negocio (usuario sin tipo de vinculación, usuario sin rol válido,
slug desconocido).
"""
import zipfile
from io import BytesIO

import pytest

from apps.sigesi.models import User
from apps.sigesi.services.formatos_docente_service import construir_contexto_formato
from apps.sigesi.views.reports.formatos_docente_view import _safe_media_path


BULK = '/api/v1/informes/formularios-docente/'
DETAIL = '/api/v1/informes/formularios-docente/descargar/'
ZIP_MAGIC = b'PK\x03\x04'  # un .zip (y un .docx) empieza con esta firma


def _body(resp):
    """Concatena el contenido de una respuesta de streaming (FileResponse)."""
    return b''.join(resp.streaming_content)


def _document_xml(resp):
    """Extrae el texto de ``word/document.xml`` del .docx devuelto en la respuesta."""
    cuerpo = _body(resp)
    assert cuerpo[:4] == ZIP_MAGIC
    with zipfile.ZipFile(BytesIO(cuerpo)) as zf:
        return zf.read('word/document.xml').decode('utf-8')


@pytest.fixture
def plantilla_tmp(tmp_path, monkeypatch):
    """Crea una plantilla .docx temporal con marcadores y la registra en la vista.

    Genera un .docx con ``{{ director_nombre }}`` y ``{{ semillero_nombre }}``,
    apunta ``FORMATOS_ROOT`` a la carpeta temporal y registra el slug
    ``plantilla-test`` para poder ejercitar la inyección de extremo a extremo.
    """
    from docx import Document
    from apps.sigesi.views.reports import formatos_docente_view as fdv

    doc = Document()
    doc.add_paragraph('Director: {{ director_nombre }}')
    doc.add_paragraph('Semillero: {{ semillero_nombre }}')
    ruta = tmp_path / 'test.docx'
    doc.save(str(ruta))

    monkeypatch.setattr(fdv, 'FORMATOS_ROOT', str(tmp_path))
    monkeypatch.setitem(fdv.FORMATOS_DOCENTE, 'plantilla-test', 'test.docx')
    return ruta


# ---------------------------------------------------------------------------
# Camino feliz — paquete .zip por tipo de vinculación
# ---------------------------------------------------------------------------

@pytest.mark.django_db
@pytest.mark.parametrize('tipo, zip_name', [
    ('catedratico', 'formatos_catedratico.zip'),
    ('planta', 'formatos_planta.zip'),
])
def test_admin_descarga_zip_segun_vinculacion(auth_client, admin_user, director_semillero, tipo, zip_name):
    director_semillero.tipo_vinculacion = tipo
    director_semillero.save()

    resp = auth_client(admin_user).get(f'{BULK}?user={director_semillero.id}')

    assert resp.status_code == 200, resp.content[:200]
    assert zip_name in resp['Content-Disposition']
    assert _body(resp)[:4] == ZIP_MAGIC


@pytest.mark.django_db
def test_director_semillero_descarga_su_propio_zip(auth_client, director_semillero):
    director_semillero.tipo_vinculacion = 'catedratico'
    director_semillero.save()

    resp = auth_client(director_semillero).get(f'{BULK}?user={director_semillero.id}')

    assert resp.status_code == 200
    assert 'formatos_catedratico.zip' in resp['Content-Disposition']


@pytest.mark.django_db
def test_admin_target_descarga_paquete_administrador(auth_client, admin_user):
    # El usuario objetivo es un administrador (sin tipo de vinculación): obtiene
    # el paquete de administrador.
    resp = auth_client(admin_user).get(f'{BULK}?user={admin_user.id}')

    assert resp.status_code == 200, resp.content[:200]
    assert 'formatos_catedratico.zip' in resp['Content-Disposition']
    assert _body(resp)[:4] == ZIP_MAGIC


# ---------------------------------------------------------------------------
# Reglas de negocio — 400 / 404
# ---------------------------------------------------------------------------

@pytest.mark.django_db
def test_400_si_usuario_sin_tipo_vinculacion(auth_client, admin_user, director_semillero):
    # director_semillero.tipo_vinculacion queda en None por defecto
    resp = auth_client(admin_user).get(f'{BULK}?user={director_semillero.id}')

    assert resp.status_code == 400
    assert 'tipo de vinculación' in resp.json()['message']


@pytest.mark.django_db
def test_400_si_usuario_no_es_director_semillero(auth_client, admin_user, estudiante):
    estudiante.tipo_vinculacion = 'planta'
    estudiante.save()

    resp = auth_client(admin_user).get(f'{BULK}?user={estudiante.id}')

    assert resp.status_code == 400
    assert 'director de semillero' in resp.json()['message']


@pytest.mark.django_db
def test_400_si_falta_parametro_user(auth_client, admin_user):
    resp = auth_client(admin_user).get(BULK)
    assert resp.status_code == 400


@pytest.mark.django_db
def test_400_si_user_no_es_entero(auth_client, admin_user):
    resp = auth_client(admin_user).get(f'{BULK}?user=abc')
    assert resp.status_code == 400


@pytest.mark.django_db
def test_404_si_usuario_no_existe(auth_client, admin_user):
    resp = auth_client(admin_user).get(f'{BULK}?user=999999')
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Permisos — 403 para roles no permitidos
# ---------------------------------------------------------------------------

@pytest.mark.django_db
@pytest.mark.parametrize('rol', ['estudiante', 'lider_estudiantil', 'director_grupo'])
def test_403_roles_no_permitidos_en_bulk(auth_client, director_semillero, rol):
    director_semillero.tipo_vinculacion = 'planta'
    director_semillero.save()
    usuario = User.objects.create(
        username=f'u_{rol}', cedula=f'CC{rol}', correo_personal=f'{rol}@x.com',
        email=f'{rol}@inst.edu', roles=[rol],
    )

    resp = auth_client(usuario).get(f'{BULK}?user={director_semillero.id}')

    assert resp.status_code == 403


# ---------------------------------------------------------------------------
# Formato individual
# ---------------------------------------------------------------------------

@pytest.mark.django_db
@pytest.mark.parametrize('slug', [
    'plan-accion-semillero', 'plan-accion-grupo', 'gestion-semillero',
    'solicitud-horas-directores', 'control-cumplimiento-produccion', 'informe-mensual',
])
def test_descarga_formato_individual(auth_client, admin_user, slug):
    resp = auth_client(admin_user).get(f'{DETAIL}?form_name={slug}')

    assert resp.status_code == 200, resp.content[:200]
    assert 'attachment' in resp['Content-Disposition']
    assert len(_body(resp)) > 0


@pytest.mark.django_db
def test_formato_individual_lo_descarga_director_semillero(auth_client, director_semillero):
    resp = auth_client(director_semillero).get(f'{DETAIL}?form_name=plan-accion-semillero')
    assert resp.status_code == 200


@pytest.mark.django_db
def test_formato_individual_slug_insensible_a_mayusculas(auth_client, admin_user):
    resp = auth_client(admin_user).get(f'{DETAIL}?form_name=PLAN-ACCION-SEMILLERO')
    assert resp.status_code == 200


@pytest.mark.django_db
def test_400_formato_individual_sin_form_name(auth_client, admin_user):
    resp = auth_client(admin_user).get(DETAIL)
    assert resp.status_code == 400


@pytest.mark.django_db
def test_404_formato_individual_desconocido(auth_client, admin_user):
    resp = auth_client(admin_user).get(f'{DETAIL}?form_name=no-existe')
    assert resp.status_code == 404


@pytest.mark.django_db
def test_403_rol_no_permitido_en_formato_individual(auth_client, estudiante):
    resp = auth_client(estudiante).get(f'{DETAIL}?form_name=plan-accion-semillero')
    assert resp.status_code == 403


# ---------------------------------------------------------------------------
# Guarda anti-traversal (defensa en profundidad de _safe_media_path)
# ---------------------------------------------------------------------------

def test_safe_media_path_bloquea_traversal():
    assert _safe_media_path('../../config/settings.py') is None
    assert _safe_media_path('no/existe.docx') is None


# ---------------------------------------------------------------------------
# Inyección de datos — servicio de contexto
# ---------------------------------------------------------------------------

@pytest.mark.django_db
def test_contexto_director_con_semillero(director_semillero, semillero_aprobado):
    director_semillero.tipo_vinculacion = 'catedratico'
    director_semillero.save()

    ctx = construir_contexto_formato(director_semillero)

    assert ctx['director_nombre'] == director_semillero.get_full_name()
    assert ctx['director_cedula'] == director_semillero.cedula
    assert ctx['director_tipo_vinculacion'] == 'Catedrático'
    assert ctx['semillero_nombre'] == 'Semillero Beta'
    assert ctx['semillero_codigo'] == 'S1'
    assert ctx['grupo_nombre'] == 'Grupo Alpha'
    assert ctx['programa_nombre'] == 'Ing. Sistemas'
    assert ctx['programa_facultad'] == 'Facultad de Ingeniería'


@pytest.mark.django_db
def test_contexto_sin_semillero_tiene_todas_las_claves(director_semillero):
    # Sin semillero dirigido: todas las claves existen y las de semillero/grupo
    # quedan en cadena vacía (Jinja nunca encuentra una variable indefinida).
    ctx = construir_contexto_formato(director_semillero)

    claves = {
        'director_nombre', 'director_cedula', 'director_tipo_vinculacion',
        'director_correo', 'director_telefono', 'semillero_nombre',
        'semillero_codigo', 'semillero_objetivo', 'semillero_lineas',
        'grupo_nombre', 'grupo_codigo', 'programa_nombre', 'programa_facultad',
        'fecha_actual',
    }
    assert claves.issubset(ctx.keys())
    assert ctx['semillero_nombre'] == ''
    assert ctx['grupo_nombre'] == ''


# ---------------------------------------------------------------------------
# Inyección de datos — extremo a extremo en /descargar/
# ---------------------------------------------------------------------------

@pytest.mark.django_db
def test_descargar_inyecta_datos_del_director(auth_client, director_semillero, semillero_aprobado, plantilla_tmp):
    director_semillero.tipo_vinculacion = 'catedratico'
    director_semillero.save()

    resp = auth_client(director_semillero).get(f'{DETAIL}?form_name=plantilla-test')

    assert resp.status_code == 200, resp.content[:200]
    xml = _document_xml(resp)
    assert director_semillero.get_full_name() in xml
    assert 'Semillero Beta' in xml
    assert '{{' not in xml  # los marcadores fueron sustituidos


@pytest.mark.django_db
def test_admin_inyecta_datos_de_otro_director(auth_client, admin_user, director_semillero, semillero_aprobado, plantilla_tmp):
    resp = auth_client(admin_user).get(
        f'{DETAIL}?form_name=plantilla-test&user={director_semillero.id}')

    assert resp.status_code == 200, resp.content[:200]
    xml = _document_xml(resp)
    assert director_semillero.get_full_name() in xml
    assert 'Semillero Beta' in xml


@pytest.mark.django_db
def test_admin_user_param_no_entero(auth_client, admin_user, plantilla_tmp):
    resp = auth_client(admin_user).get(f'{DETAIL}?form_name=plantilla-test&user=abc')
    assert resp.status_code == 400


@pytest.mark.django_db
def test_admin_user_inexistente(auth_client, admin_user, plantilla_tmp):
    resp = auth_client(admin_user).get(f'{DETAIL}?form_name=plantilla-test&user=999999')
    assert resp.status_code == 404


@pytest.mark.django_db
def test_director_ignora_user_param(auth_client, director_semillero, semillero_aprobado, otro_estudiante, plantilla_tmp):
    # Un no-administrador no puede apuntar a otro usuario: recibe SUS propios datos.
    resp = auth_client(director_semillero).get(
        f'{DETAIL}?form_name=plantilla-test&user={otro_estudiante.id}')

    assert resp.status_code == 200
    xml = _document_xml(resp)
    assert director_semillero.get_full_name() in xml


@pytest.mark.django_db
def test_plantilla_sin_marcadores_se_descarga_intacta(auth_client, admin_user):
    # Las plantillas reales aún no tienen marcadores: deben seguir descargándose
    # como un .docx válido (regresión del renderizado docxtpl).
    resp = auth_client(admin_user).get(f'{DETAIL}?form_name=plan-accion-semillero')
    assert resp.status_code == 200, resp.content[:200]
    assert _body(resp)[:4] == ZIP_MAGIC
